import docx

class Interface:
    Num = ""
    Ip = ""
    Trunk = False

class device:
    type = ""
    Dev_Int = []


def name_func():
    try:
        Lab_name = input("What is the Lab file name? ")
        if not '.docx' in Lab_name:
            Lab_name = Lab_name + '.docx'
        Lab_Doc = docx.Document(Lab_name)
    except:
        print('File does not exist')
        name_func()

    Lab = []
    Table = []
    Tab_doc = Lab_Doc.tables

    for x in range(len(Lab_Doc.paragraphs)):
        Lab.append(Lab_Doc.paragraphs[x].text)

    for table in Tab_doc:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    Table.append(paragraph.text)
        Table.append('End')
    return [Lab, Table] 

Lab = name_func()
Text = Lab[0]
Table = Lab[1]

def Find_Question():
    print()
    
# this is when I realized this is going to be way to hard to do because all the labs are diffrent and parsing them would be almost impossible
File = open('text.txt', 'w')
for x in range(len(Text)):
    File.write(Text[x] + '\n')
File.close()
File = open('tables.txt', 'w')
for x in range(len(Table)):
    File.write(Table[x] + '\n')
File.close()